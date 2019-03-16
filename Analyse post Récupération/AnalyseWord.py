import docx
import sys
import os 
if not os.name == 'nt':
	import magic
if os.name == 'nt':
	print("#######################################################################################################################\n Be careful, on windows there will be some false positive (empty result file) du to a library not available on windows \n#######################################################################################################################")
if len(sys.argv) >=2 :#test si le parametre du dossier est bien donne
	countError=0;
	for path, dirs, files in os.walk(sys.argv[1]):#parcours de dossier et sous dossier
		for filename in files:
			extension=filename.split('.')[-1] 
			if extension == "docx" :
				chemincomplet = os.path.join(path , filename)
				statinfo = os.stat(chemincomplet)
				if statinfo.st_size > 0 :
					try : # plusieur fichiers sont incomplet donc possiblement impossible a ouvrir
						if os.name == 'nt':
							document = docx.Document(chemincomplet)
							#le script va prend uniquement les fichier qui ont moins de 4 paragraphes et copie leur contenue dans un fichier text 
							if len(document.paragraphs) < 4 :
								if len(document.paragraphs[0].text) < 41 :
									if not os.path.isdir(os.path.join(path , "potentielMDP")) :
										os.mkdir(os.path.join(path , "potentielMDP"))
									fichier = open(os.path.join(path , "potentielMDP" , (filename + ".txt")),"a")
									fichier.write(document.paragraphs[0].text)
									if len(document.paragraphs) > 1 :
										fichier.write("\n"+document.paragraphs[1].text)
										if len(document.paragraphs) > 2 :
											fichier.write("\n"+document.paragraphs[2].text)
						else :
							if magic.from_file(chemincomplet) == "Microsoft Word 2007+" :  #vérifie au cas ou que le fihier est bien un docx ce qui élimine certain fichier incomplet mais fonctionne uniquement sous linux
								document = docx.Document(chemincomplet)
								#le script va prend uniquement les fichier qui ont moins de 4 paragraphes et copie leur contenue dans un fichier text 
								if len(document.paragraphs) < 4 :
									if len(document.paragraphs[0].text) < 41 :
										if not os.path.isdir(os.path.join(path , "potentielMDP")) :
											os.mkdir(os.path.join(path , "potentielMDP"))
										fichier = open(os.path.join(path , "potentielMDP" , (filename + ".txt")),"a")
										fichier.write(document.paragraphs[0].text)
										if len(document.paragraphs) > 1 :
											fichier.write("\n"+document.paragraphs[1].text)
											if len(document.paragraphs) > 2 :
												fichier.write("\n"+document.paragraphs[2].text)
					except :
						#print ("error file : "+chemincomplet) #erreur de fichier qui ne s'ouvre pas
						countError=countError+1
	print(str(countError) + " files not open du to errors")					
else : 
	print('\033[2;31;40m please give in parameter the directory to analyse.')
	print('\033[2;31;40m veuillez donner en parametre le dossier a analyser.')
